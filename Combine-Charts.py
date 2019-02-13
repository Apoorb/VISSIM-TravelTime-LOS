# -*- coding: utf-8 -*-
"""
Created on Wed Feb 13 12:35:26 2019

@author: A-Bibeka
"""

from docx import Document
from docx.shared import Inches
import os


os.chdir("D:/Dropbox/TTI_Projects/Road User Cost")

document = Document()
section = document.sections[0]
section.top_margin  = Inches(0.5)
section.bottom_margin  = Inches(0.5)


p = document.add_paragraph()
r = p.add_run()
r.add_picture('NB69_2025.png',width=Inches(4),height=Inches(2.5))
r.add_picture('NB_Spur_2025.png',width=Inches(4),height=Inches(2.5))
r.add_picture('NB69_2045.png',width=Inches(4),height=Inches(2.5))
r.add_picture('NB_Spur_2045.png',width=Inches(4),height=Inches(2.5))
r.add_picture('SB69_2025.png',width=Inches(4),height=Inches(2.5))
r.add_picture('SB_Spur_2025.png',width=Inches(4),height=Inches(2.5))
r.add_picture('SB69_2045.png',width=Inches(4),height=Inches(2.5))
r.add_picture('SB_Spur_2045.png',width=Inches(4),height=Inches(2.5))

r.add_picture('Del_NB69_2025.png',width=Inches(4),height=Inches(2.5))
r.add_picture('Del_NB_Spur_2025.png',width=Inches(4),height=Inches(2.5))
r.add_picture('Del_NB69_2045.png',width=Inches(4),height=Inches(2.5))
r.add_picture('Del_NB_Spur_2045.png',width=Inches(4),height=Inches(2.5))
r.add_picture('Del_SB69_2025.png',width=Inches(4),height=Inches(2.5))
r.add_picture('Del_SB_Spur_2025.png',width=Inches(4),height=Inches(2.5))
r.add_picture('Del_SB69_2045.png',width=Inches(4),height=Inches(2.5))
r.add_picture('Del_SB_Spur_2045.png',width=Inches(4),height=Inches(2.5))
document.save('Feb13-Charts.docx')